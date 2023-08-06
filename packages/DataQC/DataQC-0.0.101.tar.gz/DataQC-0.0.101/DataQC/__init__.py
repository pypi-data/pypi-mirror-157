################
### Packages ###
################
import os, shutil
import pandas as pd
import numpy as np
import time
import docx
from docx import Document
from docx.shared import Cm



def delete_files_from_a_folder(LOCATION):
    # Using this function we can delete all the file of a specific location in local computer.

    # Where,
    # LOCATION = Any local directory path

    folder = LOCATION
    for filename in os.listdir(folder): # This for loop sequentially delete the specific file for the location.
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))


def listToString_seperated_by_comma(s):
    # Convert a list items to string where the items seperated by comma
    # Where,
    # s = List

    str1 = ", "
    # return string
    return (str1.join(s))

def listToString_seperated_by_space(s):
    # Convert a list items to string where the items seperated by space
    # Where,
    # s = List

    str1 = ""
    # return string
    return (str1.join(s))

def combine_word_documents(files, OUTPUT_LOCATION):
    # This function combine all the specified word files to a single document and
    # save them in a specific directory

    #Where,
    # files = The word file names which are going to be merged
    # OUTPUT_LOCATION = The location where the merged document will be saved\

    merged_document = Document()
    for index, file in enumerate(files): # It combine each word file to build a single word document
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(files) - 1:
            sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save(OUTPUT_LOCATION + "/" + 'Data QC summary.docx')

def saving_data_file_names_word_local(INPUT_LOCATION, OUTPUT_LOCATION):
    # This function get all the file names from an input location and put the file names in a word file and
    # save the word file in a output location

    #Where,
    # INPUT_LOCATION = The local data directory
    # OUTPUT_LOCATION = Location where the word file will be saved

    os.chdir(INPUT_LOCATION)
    files = os.listdir()

    all_file_name = []
    for LOCALFILENAME in files: # This for loop used to append all the data file names to a single list
        all_file_name.append(LOCALFILENAME + "\n")
        all_file_name

    file_name = OUTPUT_LOCATION + "/" + 'data_file_names.docx'
    doc = Document()
    doc.add_heading("File Names:", 1)
    doc.add_paragraph(listToString_seperated_by_space(list(all_file_name)), style='List Bullet')
    doc.save(file_name)

def empty_file_checker(INPUT_LOCATION,OUTPUT_LOCATION,SEP):
    # This function helps us to find a data set or file which do not have a single line of data
    # or the file is empty

    # Where,
    # INPUT_LOCATION = The location where the data files are kept in local directory
    # OUTPUT_LOCATION = The empty file word report saving location
    # SEP = The separator used in the data (e.g. , or ' ')

    os.chdir(INPUT_LOCATION)
    files = os.listdir()

    # Selecting the data files which are less than 1 MB
    tiny_files = []
    for file in files: # This for loop create a data list where the data is less than 1 MB
        filesize = os.path.getsize(file)
        if filesize < 1000000: #Checking all the data file which are less than 1 MB
            tiny_files.append(file)

    # Selecting the empty data file from the files which are less than 1 MB
    empty_files = []
    for tiny_file_name in tiny_files: # This for loop create a data list for empty data files.
        tiny_file_path = INPUT_LOCATION + '/' + tiny_file_name
        data = pd.read_csv(tiny_file_path, sep = SEP, encoding='windows-1252', error_bad_lines=False)
        if len(data) == 0:
            empty_files.append(tiny_file_name + "\n")

    file_name = OUTPUT_LOCATION + "/" + 'empty_file_names.docx'
    doc = Document()
    doc.add_heading("Empty File Names:", 1)
    if len(empty_files) == 0:
        doc.add_paragraph("No empty file!")
    else:
        doc.add_paragraph(listToString_seperated_by_space(list(empty_files)), style='List Bullet')
    doc.save(file_name)


def checking_missing_file(LOCATION):
    # This function use to get the file names of a specific directory

    # Where,
    # LOCATION = The local directory location where the files are kept

    os.chdir(LOCATION)
    data_files = os.listdir()
    data_file_name = []

    for s in data_files: # This for loop help to create a file name list using different data names
        result = ''.join([i for i in s if not i.isdigit()])
        data_file_name.append(result)
        data_file_name
    return data_file_name


def get_data_and_column_names(LOCATION, SEP):
    # This function use to create a pandas data frame where the data file name and column names are kept.
    # The pandas data frame entries are, 1. The data file names 2. The column names of that specific data file.
    # This outputs would be use to check the data structure with the reference file.

    # Where,
    # LOCATION = The local directory location where the files are kept
    # SEP = The separator used in the data


    os.chdir(LOCATION)
    files = os.listdir()

    data_file_names = checking_missing_file(LOCATION)

    column_names = []
    for LOCALFILENAME in files: # Using this for loop we can get the column names for a specific data
        file_path = LOCATION + '/' + LOCALFILENAME
        data = pd.read_csv(file_path, sep=SEP, encoding='windows-1252', error_bad_lines=False)
        columns = data.columns
        column_names.append(columns)
        column_names

    df = pd.DataFrame(list(zip(data_file_names, column_names)),
                      columns=['Data_file_name', 'Columns'])
    return df


def missing_value_for_pandas_data(df):
    count_missing = df.isnull().sum()
    missing_value_df = pd.DataFrame({'columns': df.columns,
                                     'missing value count': count_missing})
    return missing_value_df



def create_and_save_qc_report_for_each_data_from_local(INPUT_LOCATION, OUTPUT_LOCATION, SEP):
    # Inside this function, it completed the required analysis for each of the data file to create a word file for each.
    # This function load the data, find the descriptive statistics, missing values for each of the column and save the
    # output in a word file for each of the data set.

    # Where,
    # INPUT_LOCATION = The data files directory.
    # OUTPUT_LOCATION = The output directory where the word report of each data file is saved.
    # ID_COLUMN = Previously known ID column by which we can get the unique count of the item.
    # NO_MEAN_STDDEV = The variable for which we are not interested to calculated the mean and standard deviation.
    # SEP = The separator used in the data.


    os.chdir(INPUT_LOCATION)
    files = os.listdir()

    for LOCALFILENAME in files: # This for loop create word file for each data set.
        file_path = INPUT_LOCATION + '/' + LOCALFILENAME

        start_time = time.time()
        data = pd.read_csv(file_path, sep=SEP, encoding='windows-1252', error_bad_lines=False)
        row, col = data.shape

        duration = (time.time() - start_time) / 60
        print(LOCALFILENAME + "--- %s Read time ---" % duration)

        start_time2 = time.time()
        ### Summary of the data
        #remove_mean_stddev_columns = NO_MEAN_STDDEV
        pandas_summary = pd.merge(summary_of_pandas_data(data),
                                  missing_value_for_pandas_data(data), on='columns')  # creating summary table
        pandas_summary['missing value %'] = (pandas_summary['missing value count'].div(row)) * 100
        pandas_summary['missing value %'] = pandas_summary['missing value %'].round(2)

        duration2 = (time.time() - start_time2) / 60
        print(LOCALFILENAME + "--- %s Summary time ---" % duration2)


        mask = pandas_summary.applymap(lambda x: x is None)
        cols = pandas_summary.columns[(mask).any()]
        for col in pandas_summary[cols]:
            pandas_summary.loc[mask[col], col] = np.nan

        pandas_summary = pandas_summary.round(2)  # 2 digit round

        number_of_column = 'Number of columns = ' + str(col) + "\n"
        number_of_row = 'Number of rows = ' + f'{row:,}' + "\n"
        file_name = OUTPUT_LOCATION + '/' + "Qc for" + LOCALFILENAME + '.docx'
        doc = docx.Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        doc.add_heading(LOCALFILENAME, 1)

        if data.shape[0] > 0:
            doc.add_paragraph(number_of_column)
            doc.add_paragraph(number_of_row)
            t = doc.add_table(pandas_summary.shape[0] + 1, pandas_summary.shape[1])
            t.style = 'Table Grid'
            # add the header rows.
            for j in range(pandas_summary.shape[-1]):
                t.cell(0, j).text = pandas_summary.columns[j]
            # add the rest of the data frame
            for i in range(pandas_summary.shape[0]):
                for j in range(pandas_summary.shape[-1]):
                    t.cell(i + 1, j).text = str(pandas_summary.values[i, j])

            doc.save(file_name)
        else:
            p = doc.add_paragraph()
            runner = p.add_run("No data in the file.")
            runner.bold = True
            runner.italic = True
            doc.save(file_name)


def combining_qc_reports_word_local(INPUT_LOCATION, OUTPUT_LOCATION):
    # This function create a combined word file

    # Where,
    # INPUT_LOCATION = The data directory
    # OUTPUT_LOCATION = The output file directory or the location where the combined report would be kept.


    os.chdir(INPUT_LOCATION)
    files = os.listdir()

    all_file = []
    all_file.append('data_file_names.docx')
    all_file.append('missing_file_and_structure.docx')
    all_file.append('empty_file_names.docx')

    for LOCALFILENAME in files: # This for loop combine all the word files for the data sets.
        all_file.append("Qc for" + LOCALFILENAME + '.docx')
        all_file
    os.chdir(OUTPUT_LOCATION)
    combine_word_documents(all_file,OUTPUT_LOCATION)

def checking_missing_files(INPUT_LOCATION, REF_LOCATION):
    # Using this function we compare the data file names with the reference files.

    # Where,
    # INPUT_LOCATION = The data directory
    # REF_LOCATION = The reference file directory

    Data = checking_missing_file(LOCATION=INPUT_LOCATION)
    Reference = checking_missing_file(LOCATION=REF_LOCATION)

    # Data file name not matched with reference:
    unmatched_data = list(set(Data) - set(Reference))
    unmatched_data_text = ','.join(map(str, unmatched_data))

    # Data file name not matched with reference:
    unmatched_ref = list(set(Reference) - set(Data))
    unmatched_ref_text = ','.join(map(str, unmatched_ref))

    if len(unmatched_data) == 0 and len(unmatched_ref) == 0:
        missing_file_report = ("No data files are missing!")
    elif len(unmatched_data) or len(unmatched_ref) != 0:
        missing_files = []
        for file_missing in unmatched_ref: # Using this for loop we combine the missing file names
            missing_files.append(file_missing + "\n")
            missing_files
        unmatched_ref_text_updated = ''.join(map(str, missing_files))
        missing_file_report = unmatched_ref_text_updated
    return missing_file_report

def data_structure_checking(INPUT_LOCATION, REF_LOCATION, SEP):
    Data = get_data_and_column_names(LOCATION=INPUT_LOCATION, SEP = SEP)
    Ref = get_data_and_column_names(LOCATION=REF_LOCATION, SEP = SEP)

    merged_data = pd.merge(Data, Ref, on='Data_file_name')
    merged_data.columns = ["file_name", "columns_in_data", "columns_in_reference"]

    merged_data['extra_columns_in_data'] = (
                merged_data['columns_in_data'].apply(set) - merged_data['columns_in_reference'].apply(set)).apply(
        list).apply(listToString_seperated_by_comma) # Create a string object of the main data or given data columns

    merged_data['extra_columns_in_reference'] = (
                merged_data['columns_in_reference'].apply(set) - merged_data['columns_in_data'].apply(set)).apply(
        list).apply(listToString_seperated_by_comma) # Create a string object of the reference data columns

    merged_data = merged_data[['file_name', 'extra_columns_in_data', 'extra_columns_in_reference']]
    merged_data.columns = ['File name', 'Extra columns in data','Extra columns in reference']

    problematic_data = merged_data[
        (merged_data["Extra columns in data"] != '') != (merged_data["Extra columns in reference"] != '')] # Using the
    # intersection rule to get whether there is any uncommon column between reference and given data set.

    if len(problematic_data) == 0: # If all the columns are same in reference and given data then the length should be 0
        # and the data structure is good. Otherwise it will report the problematic column names with data file name.
        data_structure = problematic_data
    else:
        problematic_data = problematic_data.apply(lambda x: x.str.strip()).replace('', np.nan)
        problematic_data = problematic_data.dropna(how='all', axis=1)
        data_structure = problematic_data

    return (data_structure)


def missing_file_and_structure_checking(INPUT_LOCATION, OUTPUT_LOCATION, REF_LOCATION, SEP):
    # This function actually combine the two function called checking_missing_files and data_structure_checking to create a
    # combined word report for them
    missing_file_report = checking_missing_files(INPUT_LOCATION, REF_LOCATION)
    data_structure_report = data_structure_checking(INPUT_LOCATION, REF_LOCATION, SEP)

    ### Saving output in word formate:
    file_name = OUTPUT_LOCATION + '/' + "missing_file_and_structure" + '.docx'
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    doc.add_heading("Missing file report:", 1)
    doc.add_paragraph(missing_file_report)
    doc.add_heading("Data structure report:", 1)
    if len(data_structure_report) == 0:
        data_structure_report = "Data structures are fine for all the data files."
        doc.add_paragraph(data_structure_report)
    else:
        df = data_structure_report
        t = doc.add_table(df.shape[0] + 2, df.shape[1])
        t.style = 'Table Grid'

        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])

    doc.save(file_name)

def summary_of_pandas_data(data):
    numerical_columns = data.describe().T.round(decimals=2)
    object_columns = data.describe(include=['O']).T
    appended_data = object_columns.append(numerical_columns)

    pandas_summary = appended_data.replace(np.nan, '', regex=True)
    pandas_summary = pandas_summary[1:]  # take the data less the header row

    pandas_summary = pandas_summary.reset_index()
    pandas_summary = pandas_summary.rename(columns={'index': 'columns'})

    pandas_summary = pandas_summary.drop_duplicates(subset=['columns'], keep='last')

    return pandas_summary


def qc_report_local_folder(INPUT_LOCATION, OUTPUT_LOCATION, REF_LOCATION,SEP):
    # This function used to combine the saving_data_file_names_word_local, missing_file_and_structure_checking,
    # empty_file_checker, create_and_save_qc_report_for_each_data_from_local, and combining_qc_reports_word_local functions!

    start_time = time.time()

    # Getting the file names
    saving_data_file_names_word_local(INPUT_LOCATION, OUTPUT_LOCATION)

    duration1 = (time.time() - start_time) / 60
    print("--- %s Saving data file names ---" % duration1)

    start_time2 = time.time()
    # Missing file and data structure:
    missing_file_and_structure_checking(INPUT_LOCATION, OUTPUT_LOCATION, REF_LOCATION, SEP)

    duration2 = (time.time() - start_time2) / 60
    print("--- %s Missing file and structure checking ---" % duration2)

    # Empty file checking
    start_time3 = time.time()
    empty_file_checker(INPUT_LOCATION, OUTPUT_LOCATION, SEP)
    duration3 = (time.time() - start_time3) / 60
    print("--- %s Empty file checking ---" % duration3)

    start_time5 = time.time()
    # Creating individual report for each data:
    create_and_save_qc_report_for_each_data_from_local(INPUT_LOCATION, OUTPUT_LOCATION, SEP)  # creating qc word report for each data and saving

    duration5 = (time.time() - start_time5) / 60
    print("--- %s Report generating for all data ---" % duration5)

    start_time6 = time.time()
    # Combining the reports:
    combining_qc_reports_word_local(INPUT_LOCATION, OUTPUT_LOCATION)
    duration6 = (time.time() - start_time6) / 60
    print("--- %s Combining report ---" % duration6)

    duration = (time.time() - start_time) / 60
    print("--- %s Total minutes ---" % duration)

def Quick_QC(INPUT_LOCATION,
        OUTPUT_LOCATION,
        REF_LOCATION,
        SEP):

    # This is the main application user will use. In this function we combined all the parts together. Here we combined
    # delete_files_from_a_folder, delete_files_from_a_folder, qc_report_local_folder and the last part added the email functionalities.

    if len(os.listdir(INPUT_LOCATION)) == 0:
        print("Location is empty!")

    else:
        # Create the reports:
        qc_report_local_folder(INPUT_LOCATION, OUTPUT_LOCATION, REF_LOCATION,SEP)

        print("Pipeline executed successfully!")






